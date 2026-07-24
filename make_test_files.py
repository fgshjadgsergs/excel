# -*- coding: utf-8 -*-
"""
Генератор тестовых файлов для надстройки SmartCells.

Создаёт:
  test/Тест_данные.xlsx   — книга с листом «Данные» и набором ячеек
                            разных форматов (число с разделителями, дата,
                            процент, длинный текст, формула, ошибка, пустая).
  test/Тест_документ.docx — документ с метками на все эти ячейки,
                            плюс кириллическая метка {А1} и заведомо
                            кривая {5A}.

Запуск:
    pip install openpyxl python-docx
    python make_test_files.py
"""

import os
import datetime

from openpyxl import Workbook
from docx import Document


OUT_DIR = "test"
XLSX_PATH = os.path.join(OUT_DIR, "Тест_данные.xlsx")
DOCX_PATH = os.path.join(OUT_DIR, "Тест_документ.docx")
SHEET_NAME = "Данные"


def make_xlsx():
    """Собрать книгу с ячейками разных типов и форматов."""
    wb = Workbook()
    ws = wb.active
    ws.title = SHEET_NAME

    # A1 — число с разделителями разрядов
    ws["A1"] = 1234567.89
    ws["A1"].number_format = "# ##0,00"

    # A2 — дата
    ws["A2"] = datetime.date(2026, 7, 22)
    ws["A2"].number_format = "DD.MM.YYYY"

    # A3 — процент
    ws["A3"] = 0.256
    ws["A3"].number_format = "0.0%"

    # A4 — длинный текст
    ws["A4"] = ("Длинный текст для проверки подстановки значения "
                "в документ Word без потери форматирования абзаца.")

    # A5 — формула (значение вычислит Excel при открытии)
    ws["A5"] = "=A1*2"
    ws["A5"].number_format = "# ##0,00"

    # A6 — ячейка с ошибкой (деление на ноль)
    ws["A6"] = "=1/0"

    # A7 — пустая ячейка (намеренно не заполняем)
    #      оставляем без значения

    # Подписи в соседней колонке — чтобы файл было удобно смотреть вручную
    labels = {
        "B1": "число с разделителями",
        "B2": "дата",
        "B3": "процент",
        "B4": "длинный текст",
        "B5": "формула =A1*2",
        "B6": "ошибка =1/0",
        "B7": "пустая ячейка",
    }
    for addr, text in labels.items():
        ws[addr] = text

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 28

    # Второй лист — для проверки меток с указанием листа {Лист2!A1}
    ws2 = wb.create_sheet("Лист2")
    ws2["A1"] = "Значение со второго листа"

    os.makedirs(OUT_DIR, exist_ok=True)
    wb.save(XLSX_PATH)
    print("Создан:", XLSX_PATH)


def make_docx():
    """Собрать документ с метками на все ячейки + ошибочные метки."""
    doc = Document()
    doc.add_heading("Тестовый документ SmartCells", level=1)

    doc.add_paragraph(
        "Ниже метки, которые надстройка должна заменить значениями "
        "из файла «Тест_данные.xlsx», лист «Данные»."
    )

    rows = [
        ("Число с разделителями", "{A1}"),
        ("Дата", "{A2}"),
        ("Процент", "{A3}"),
        ("Длинный текст", "{A4}"),
        ("Формула", "{A5}"),
        ("Ячейка с ошибкой", "{A6}"),
        ("Пустая ячейка", "{A7}"),
    ]
    for name, mark in rows:
        doc.add_paragraph("{}: {}".format(name, mark))

    # Метка с указанием другого листа
    doc.add_paragraph("Значение с другого листа: {Лист2!A1}")

    doc.add_paragraph("")
    doc.add_paragraph("Ошибочные метки (должны попасть в список проблем):")

    # Кириллическая «А» вместо латинской — метка {А1}
    doc.add_paragraph("Кириллическая метка: {А1}")
    # Заведомо кривой адрес — цифра перед буквой
    doc.add_paragraph("Кривая метка: {5A}")

    os.makedirs(OUT_DIR, exist_ok=True)
    doc.save(DOCX_PATH)
    print("Создан:", DOCX_PATH)


if __name__ == "__main__":
    make_xlsx()
    make_docx()
    print("Готово. Проверьте папку:", OUT_DIR)
